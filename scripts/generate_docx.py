#!/usr/bin/env python3
"""Generate a Chinese official-document style .docx from controlled Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from font_guard import build_font_config, ensure_required_fonts, install_assets_for_missing

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - user environment dependent
    raise SystemExit(
        "Missing dependency: python-docx. Install it with `pip install python-docx`."
    ) from exc


FRONT_RE = re.compile(r"\A---\s*\n(?P<body>.*?)\n---\s*\n?", re.S)


def parse_front_matter(text: str) -> tuple[dict[str, object], str]:
    match = FRONT_RE.match(text)
    if not match:
        return {}, text

    meta: dict[str, object] = {}
    current_list: list[str] | None = None
    current_key: str | None = None

    for raw in match.group("body").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.lstrip().startswith("- ") and current_list is not None:
            current_list.append(line.split("- ", 1)[1].strip())
            continue
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            current_key = key
            if value:
                meta[key] = value
                current_list = None
            else:
                current_list = []
                meta[key] = current_list
        elif current_key:
            meta[current_key] = line.strip()

    return meta, text[match.end() :]


def set_run_font(
    run,
    font_name: str,
    size_pt: float | None = None,
    bold: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = font_name
    if size_pt:
        run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        normalized = color.lstrip("#")
        run.font.color.rgb = RGBColor(
            int(normalized[0:2], 16),
            int(normalized[2:4], 16),
            int(normalized[4:6], 16),
        )
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    font: str = "仿宋_GB2312",
    size: float = 16,
    bold: bool = False,
    color: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent: bool = True,
    line_spacing: float = 28,
):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(line_spacing)
    if first_indent:
        fmt.first_line_indent = Pt(32)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold, color)
    return paragraph


def add_bottom_border(paragraph, *, color: str = "FF0000", size: int = 12, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_red_separator(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(14)
    paragraph.paragraph_format.line_spacing = Pt(4)
    add_bottom_border(paragraph)


def add_issue_line(
    doc: Document,
    *,
    doc_number: str = "",
    issue_person: str = "",
    fonts: dict[str, str],
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if doc_number else WD_ALIGN_PARAGRAPH.RIGHT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(28)

    if doc_number:
        run = paragraph.add_run(doc_number)
        set_run_font(run, fonts["body"], 16)
    if doc_number and issue_person:
        spacer = paragraph.add_run("        ")
        set_run_font(spacer, fonts["body"], 16)
    if issue_person:
        label = paragraph.add_run("签发人：")
        set_run_font(label, fonts["body"], 16)
        name = paragraph.add_run(issue_person)
        set_run_font(name, fonts["h2"], 16)


def add_page_number_to_footer(footer, fonts: dict[str, str], align: WD_ALIGN_PARAGRAPH) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run("— ")
    set_run_font(run, fonts["footer"], 14)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    field_run._r.append(fld_begin)
    field_run._r.append(instr_text)
    field_run._r.append(fld_end)
    set_run_font(field_run, fonts["footer"], 14)

    run = paragraph.add_run(" —")
    set_run_font(run, fonts["footer"], 14)


def add_page_numbers(
    doc: Document,
    section,
    fonts: dict[str, str],
    *,
    show_first_page_number: bool = False,
    page_number_position: str = "outer",
) -> None:
    section.different_first_page_header_footer = not show_first_page_number

    if page_number_position == "center":
        add_page_number_to_footer(section.footer, fonts, WD_ALIGN_PARAGRAPH.CENTER)
        if show_first_page_number:
            add_page_number_to_footer(section.first_page_footer, fonts, WD_ALIGN_PARAGRAPH.CENTER)
        return

    doc.settings.odd_and_even_pages_header_footer = True
    add_page_number_to_footer(section.footer, fonts, WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number_to_footer(section.even_page_footer, fonts, WD_ALIGN_PARAGRAPH.LEFT)
    if show_first_page_number:
        add_page_number_to_footer(section.first_page_footer, fonts, WD_ALIGN_PARAGRAPH.RIGHT)


def setup_document(
    fonts: dict[str, str],
    *,
    show_first_page_number: bool = False,
    page_number_position: str = "outer",
) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(28)
    add_page_numbers(
        doc,
        section,
        fonts,
        show_first_page_number=show_first_page_number,
        page_number_position=page_number_position,
    )

    style = doc.styles["Normal"]
    style.font.name = fonts["body"]
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), fonts["body"])
    return doc


def meta_text(meta: dict[str, object], key: str) -> str:
    value = meta.get(key, "")
    if isinstance(value, list):
        return "；".join(str(item).strip() for item in value if str(item).strip())
    return str(value).strip()


def truthy(value: object, *, default: bool = False) -> bool:
    if value is None or str(value).strip() == "":
        return default
    normalized = str(value).strip().lower()
    return normalized not in {"0", "false", "no", "off", "none", "否", "不", "无"}


def add_document_header(doc: Document, meta: dict[str, object], fonts: dict[str, str]) -> None:
    issuer_mark = meta_text(meta, "issuer_mark")
    doc_number = meta_text(meta, "doc_number")
    issue_person = (
        meta_text(meta, "issue_person")
        or meta_text(meta, "signatory")
        or meta_text(meta, "issuer_person")
    )
    preface_items = [
        meta_text(meta, "copy_number"),
        meta_text(meta, "secret_level"),
        meta_text(meta, "urgency"),
    ]
    has_header = any([issuer_mark, doc_number, issue_person, *preface_items])
    if not has_header:
        return

    red_head = truthy(meta.get("red_head"), default=True)

    for item in preface_items:
        if item:
            add_paragraph(
                doc,
                item,
                font=fonts["body"],
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_indent=False,
                line_spacing=28,
            )

    if issuer_mark:
        add_paragraph(
            doc,
            issuer_mark,
            font=fonts["issuer"],
            size=28,
            color="FF0000" if red_head else None,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=False,
            line_spacing=34,
        )

    if doc_number and issue_person:
        add_issue_line(doc, doc_number=doc_number, issue_person=issue_person, fonts=fonts)
    elif doc_number:
        add_paragraph(doc, doc_number, font=fonts["body"], align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    elif issue_person:
        add_issue_line(doc, issue_person=issue_person, fonts=fonts)

    if red_head:
        add_red_separator(doc)


def clean_heading(line: str, level: int) -> str:
    return line[level:].strip()


def add_markdown_lines(doc: Document, lines: list[str], fonts: dict[str, str]) -> None:
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            add_paragraph(
                doc,
                clean_heading(line, 1),
                font=fonts["title"],
                size=22,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_indent=False,
                line_spacing=30,
            )
        elif line.startswith("## "):
            add_paragraph(doc, clean_heading(line, 2), font=fonts["h1"], size=16, bold=False, first_indent=False)
        elif line.startswith("### "):
            add_paragraph(doc, clean_heading(line, 3), font=fonts["h2"], size=16, first_indent=True)
        elif line.startswith("#### "):
            add_paragraph(doc, clean_heading(line, 4), font=fonts["h3"], size=16, bold=True, first_indent=True)
        else:
            no_indent = line.endswith("：") or line.startswith(("附件：", "出席：", "列席："))
            add_paragraph(doc, line, font=fonts["body"], first_indent=not no_indent)


def append_metadata(doc: Document, meta: dict[str, object], fonts: dict[str, str]) -> None:
    signer = str(meta.get("signer", "")).strip()
    date = str(meta.get("date", "")).strip()
    attachments = meta.get("attachments")

    if attachments:
        values = attachments if isinstance(attachments, list) else [str(attachments)]
        add_paragraph(doc, "附件：" + "；".join(str(v) for v in values), font=fonts["body"], first_indent=False)

    if signer:
        add_paragraph(doc, signer, font=fonts["body"], align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    if date:
        add_paragraph(doc, date, font=fonts["body"], align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)


def unique_output_path(output_path: Path) -> Path:
    if not output_path.exists():
        return output_path

    stem = output_path.stem
    suffix = output_path.suffix
    parent = output_path.parent
    for idx in range(2, 1000):
        candidate = parent / f"{stem}-v{idx:02d}{suffix}"
        if not candidate.exists():
            return candidate
    raise SystemExit(f"Too many existing versions for {output_path}")


def build_docx(
    input_path: Path,
    output_path: Path,
    *,
    install_font_assets: bool = False,
    overwrite: bool = False,
) -> Path:
    text = input_path.read_text(encoding="utf-8-sig")
    meta, body = parse_front_matter(text)
    lines = body.splitlines()

    fonts = build_font_config(meta)
    if install_font_assets:
        install_assets_for_missing(fonts)
    ensure_required_fonts(fonts)

    show_first_page_number = truthy(meta.get("show_first_page_number"), default=False)
    page_number_position = meta_text(meta, "page_number_position") or "outer"
    if page_number_position not in {"outer", "center"}:
        raise SystemExit("page_number_position must be `outer` or `center`.")

    doc = setup_document(
        fonts,
        show_first_page_number=show_first_page_number,
        page_number_position=page_number_position,
    )
    recipients = str(meta.get("recipients", "")).strip()

    add_document_header(doc, meta, fonts)

    if recipients and recipients not in body:
        recipient_line = recipients if recipients.endswith("：") else f"{recipients}："
        insert_at = 0
        for idx, line in enumerate(lines):
            if line.strip().startswith("# "):
                insert_at = idx + 1
                break
        lines.insert(insert_at, recipient_line)

    add_markdown_lines(doc, lines, fonts)

    append_metadata(doc, meta, fonts)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    final_output = output_path if overwrite else unique_output_path(output_path)
    doc.save(final_output)
    return final_output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", help="UTF-8 controlled Markdown file")
    parser.add_argument("-o", "--output", required=True, help="Output .docx path")
    parser.add_argument("--doc-type", default="", help="Document type, used for naming/logging")
    parser.add_argument(
        "--install-font-assets",
        action="store_true",
        help="Install matching authorized font files from assets/fonts before export",
    )
    parser.add_argument("--overwrite", action="store_true", help="Overwrite output path instead of creating a -vNN file")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")
    final_output = build_docx(
        input_path,
        output_path,
        install_font_assets=args.install_font_assets,
        overwrite=args.overwrite,
    )
    print(f"Generated: {final_output.resolve()}")


if __name__ == "__main__":
    main()
